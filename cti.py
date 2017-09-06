#!/usr/bin/python3
# Cyber Threat Information

from docx import Document
from sys import argv
from isac import ipcountry, insert, con, cur
from re import compile

doc = Document(argv[1])

# "IP Address"라는 문자열이 있는 테이블 번호와
# 칼럼의 번호를 파악하고 리스트에 튜플로 저장한다.
# "Attack Type"도 같은 원리로 동작.
# 다른 칼럼도 마찬가지
ipindex = []
attacktypeindex= []
urlindex = []
targetindex = []
for i, table in enumerate(doc.tables):
	for j, head in enumerate(table.rows[0].cells):
		if "IP Address" in head.text:
			ipindex.append((i, j))
		if "Attack Type" in head.text:
			attacktypeindex.append((i, j))
		if "URL" in head.text:
			urlindex.append((i, j))
		if "Target" in head.text:
			targetindex.append((i, j))

# 테이블과 칼럼 번호가 적혀 있는 인덱스를 튜플로 받고
# 칼럼 이름을 인자로 넣어주면
# 테이블에서 해당 칼럼의 데이터만 모아서 리스트로 리턴
def parsecol(index, name):
	t = []
	for table, col in index:
		cells = doc.tables[table].columns[col].cells
		for cell in cells:
			if name not in cell.text:
				t.append((cell.text,))
	return t

# URL에서 도메인만 파싱하는 정규식
def gethost(url):
	p = compile('hxxps{0,1}://(.*?)/')
	m = p.search(url[0])
	return (m.group(1),)

# 아이피만 추출해서 국가를 조회하고
# 데이터베이스에 입력한다.
#iplist = parsecol(ipindex, "IP Address")
#ip = ipcountry(iplist)
#insert(ip, 'ip')

# 공격유형을 중복제거하고 화면에 출력
#attacktypelist = parsecol(attacktypeindex, "Attack Type")
#for i in set(attacktypelist):
#	print(i)

# 호스트를 데이터베이스에 입력
#urllist = parsecol(urlindex, "URL")
#host = list(map(gethost, urllist))
#insert(host, 'host')

# 타겟을 화면에 출력
#targetlist = parsecol(targetindex, "Target")
#for i in set(targetlist): print(i)

# 문서의 날짜를 찾는 코드
date = doc.tables[0].rows[2].cells[1].text
p = compile('(\d{4})년 (\d{2})월 (\d{2})일')
m= p.search(date)
date = '-'.join([m.group(1), m.group(2), m.group(3)])

threat = []
urllist = []

for table in doc.tables:
	for cell in table.rows[0].cells:
		# 'Level' 칼럼은 사이버 위협 정보 테이블에만 존재한다.
		if 'Level' in cell.text:
			for row in table.rows[1:]:
				ip = row.cells[1].text
				attacktype = row.cells[3].text
				level = row.cells[4].text
				# 공격유형 이름을 아이디로 변환한다.
				cur.execute('select attacktypeid from attacktype where attacktypename="%s"' % attacktype)
				attacktype = cur.fetchone()[0]
				# 레벨 이름을 레벨 아이디로 변환한다.
				cur.execute('select levelid from level where levelname="%s"' % level)
				level = cur.fetchone()[0]
				t = (date, ip, attacktype, 2, level)
				threat.append(t)
		# 'Target' 칼럼은 피싱사이트 테이블에만 존재한다.
		if 'Target' in cell.text:
			for row in table.rows[1:]:
				url = row.cells[0].text
				ip = row.cells[1].text
				target = row.cells[3].text
				t = (url, ip, 4, 2, 4)
				urllist.append(t)
	# 'RFI' 테이블은 칼럼의 개수가 3개이다.
	if len(table.rows[0].cells) == 3:
		for row in table.rows[1:]:
			url = row.cells[0].text
			ip = row.cells[1].text
			t = (url, ip, 6, 2, 1)
			urllist.append(t)

# 사이버 위협 정보를 데이터베이스에 입력한다.
#insert(threat, '"ip/attacktype/src"')
# URL을 데이터베이스에 입력한다.
# URL은 피싱사이트와 RFI를 합친 것이다.
insert(urllist, 'url')

con.commit()
con.close()
